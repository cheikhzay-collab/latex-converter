import html
import io
import re
import markdown
import latex2mathml.converter
from flask import Flask, render_template, request, send_file, make_response, Response
from functools import wraps
import sqlite3
import uuid
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from lxml import etree
import lxml.html

app = Flask(__name__)

if os.environ.get('VERCEL_ENV') or os.environ.get('VERCEL'):
    DB_PATH = '/tmp/stats.db'
else:
    DB_PATH = 'stats.db'

def init_db():
    try:
        with sqlite3.connect(DB_PATH) as conn:
            c = conn.cursor()
            c.execute('''CREATE TABLE IF NOT EXISTS stats (id INTEGER PRIMARY KEY, downloads INTEGER, chars INTEGER, unique_users INTEGER)''')
            c.execute('''CREATE TABLE IF NOT EXISTS daily_stats (date TEXT PRIMARY KEY, downloads INTEGER DEFAULT 0, chars INTEGER DEFAULT 0, unique_users INTEGER DEFAULT 0)''')
            c.execute('''CREATE TABLE IF NOT EXISTS users (visitor_id TEXT PRIMARY KEY, first_visit_date TEXT)''')
            c.execute('''CREATE TABLE IF NOT EXISTS conversions (
                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                visitor_id TEXT,
                content TEXT,
                char_count INTEGER,
                is_rtl INTEGER
            )''')
            conn.commit()
    except Exception as e:
        print(f"Failed to initialize database: {e}")

init_db()

# Use local copy of MML2OMML.XSL so it works in production too
XSLT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "MML2OMML.XSL")

def clean_raw_brackets(text):
    """Fixes [ ... ] blocks which frequently appear when copying from ChatGPT or Gemini directly."""
    # Normalize non-breaking spaces before matching
    text = text.replace('\xa0', ' ')
    
    def bracket_math_cleaner(match):
        content = match.group(1).strip()
        # French interval check: [digit, digit[ or [digit; digit]
        if re.search(r'^[\d., \-+]+[ \t]*[,;][ \t]*[\d., \-+\\infty]+[\[\]]?$', content):
            return match.group(0) # Keep as is
        # Strong math indicator
        if re.search(r'[\^\\=_]|\d+[a-z]|\\frac|\\int|\\sum|\\lim', content) or len(content) > 15:
            return f"$${content}$$"
        return match.group(0)

    # single-line bracket
    text = re.sub(r'(?m)^[^\S\r\n]*\[([^\[\]\n]+)\][^\S\r\n]*$', bracket_math_cleaner, text)
    # multi-line bracket
    text = re.sub(r'(?m)^[^\S\r\n]*\[\s*\n([\s\S]*?)\n[^\S\r\n]*\][^\S\r\n]*$', bracket_math_cleaner, text)
    return text

def safe_paren_math_cleaner(text):
    """
    Applies paren_math_cleaner to convert bare (E = mc^2) to $E = mc^2$,
    but safely ignores text that is already inside a LaTeX math block.
    """
    hidden = {}
    counter = 0

    def hide(match):
        nonlocal counter
        key = f"__MATH_{counter}__"
        hidden[key] = match.group(0)
        counter += 1
        return key

    # Hide existing math blocks so we don't double-process or corrupt them
    text = re.sub(r'\$\$[\s\S]*?\$\$', hide, text)      # Display math
    text = re.sub(r'\\\[[\s\S]*?\\\]', hide, text)      # LaTeX Display
    text = re.sub(r'\\\([\s\S]*?\\\)', hide, text)      # LaTeX Inline
    text = re.sub(r'\$[\s\S]*?\$', hide, text)          # MathJax Inline
    
    def paren_math_cleaner_fn(match):
        content = match.group(1).strip()
        if re.search(r'[\^\\=_]', content) and len(content) > 3:
            return f"${content}$"
        return match.group(0)

    # 1-level nested like (f'(x) = 2x)
    text = re.sub(r'(?<!\\)\(((?:[^()]|\([^()]*\))*)\)', paren_math_cleaner_fn, text)
    
    # Restore hidden math
    for key, val in hidden.items():
        text = text.replace(key, val)
        
    return text

def preprocess_chatgpt(text):
    """Safe pipeline for ChatGPT"""
    text = clean_raw_brackets(text)
    # 1. Standard LaTeX display: \[ ... \] -> $$ ... $$
    text = re.sub(r'\\\[([\s\S]*?)\\\]', r'\n\n$$\1$$\n\n', text)
    # 2. Double parentheses ((C_f)) -> \(C_f\)
    text = re.sub(r'\(\(([\s\S]*?)\)\)', r'\(\1\)', text)
    # 3. Standard LaTeX inline: \( ... \) -> $ ... $
    text = re.sub(r'\\\(([\s\S]*?)\\\)', r'$\1$', text)
    # 4. ChatGPT interval notation: ([a, b]) -> $[a, b]$
    text = re.sub(r'\(\[([^\]]*?[,;][^\]]*?)\]\)', r'$[\1]$', text)
    
    # 5. Simple parenthesis inline (E = mc^2) (Safe from existing math)
    text = safe_paren_math_cleaner(text)
    
    return text

def preprocess_gemini(text):
    """Aggressive heuristic pipeline for Gemini which has varied outputs"""
    text = clean_raw_brackets(text)
    # 1. Standard LaTeX display: \[ ... \] -> $$ ... $$
    text = re.sub(r'\\\[([\s\S]*?)\\\]', r'\n\n$$\1$$\n\n', text)
    # 2. Interval notation: ([a, b]) -> $[a, b]$
    text = re.sub(r'\(\[([^\]]*?[,;][^\]]*?)\]\)', r'$[\1]$', text)
    
    # 3. Simple parenthesis inline (E = mc^2) (Safe from existing math)
    text = safe_paren_math_cleaner(text)
    
    # 4. Double parentheses ((C_f)) -> \(C_f\)
    text = re.sub(r'\(\(([\s\S]*?)\)\)', r'\(\1\)', text)
    return text

def preprocess_copied_math(text):
    """
    Cleans up common LLM math notation and fixes the 'Interval Bug'.
    Uses Smart AI Detection to route to the correct pipeline.
    """
    # Normalize Windows line endings (\r\n -> \n) to prevent regex failures
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove AI citation markers common to both
    text = re.sub(r'\[cite[_:][^\]]*\]', '', text)
    
    # Count indicators
    chatgpt_indicators = len(re.findall(r'\\\[|\\\(', text))
    gemini_indicators = len(re.findall(r'\$\$', text))
    
    if chatgpt_indicators > 0 and chatgpt_indicators > gemini_indicators:
        print("[System] Detected: ChatGPT Source")
        return preprocess_chatgpt(text)
    else:
        print("[System] Detected: Gemini Source / Default")
        return preprocess_gemini(text)

# Cache the XSLT transformer to avoid Disk/Memory issues
try:
    _xslt_tree = etree.parse(XSLT_PATH)
    MML2OMML_TRANSFORM = etree.XSLT(_xslt_tree)
except Exception as e:
    print(f"Failed to load MML2OMML.XSL: {e}")
    MML2OMML_TRANSFORM = None

def get_omml(latex):
    """Converts LaTeX to OMML XML element using MML2OMML.XSL."""
    if MML2OMML_TRANSFORM is None:
        return None, "XSLT Transformer not loaded"
    try:
        mathml = latex2mathml.converter.convert(latex)
        # Parse MathML
        mml_root = etree.fromstring(mathml.encode('utf-8'))
        
        # Transform to OMML
        omml_tree = MML2OMML_TRANSFORM(mml_root)
        # Convert to string for python-docx to parse
        omml_str = etree.tostring(omml_tree)
        return parse_xml(omml_str), None
    except Exception as e:
        error_msg = f"{type(e).__name__}: {str(e)}"
        print(f"OMML Conversion Error for '{latex}': {error_msg}")
        return None, error_msg

def add_math_to_run(paragraph, latex, is_display=False):
    """Inserts a native Math object into a paragraph."""
    if is_display:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    omml_node, err = get_omml(latex)
    if omml_node is not None:
        paragraph._element.append(omml_node)
    else:
        # Fallback to red text if conversion fails
        run = paragraph.add_run(f"[Math Error: {latex} | {err}]")
        run.font.color.rgb = RGBColor(255, 0, 0)

def html_to_docx(html_content, document, is_rtl=False):
    """
    Maps HTML elements (from markdown) to python-docx elements.
    Handles math placeholders ($$ and $).
    """
    if not html_content.strip():
        return
        
    try:
        # We wrap in a div and use fromstring to ensure all fragments are captured
        fragment = lxml.html.fromstring(f"<div>{html_content}</div>")
        elements = fragment.xpath('./*') # Get direct children (p, h1, ul, etc.)
    except Exception as e:
        print(f"HTML Parse Error: {e}")
        return

    if not elements:
        print("No elements found in HTML content")
        return
    
    for element in elements:
        tag = element.tag
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            p = document.add_heading('', level=level)
            if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            process_inline(element, p)
            
        elif tag == 'p':
            p = document.add_paragraph()
            if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            process_inline(element, p)
            
        elif tag in ['ul', 'ol']:
            for li in element.xpath('.//li'):
                p = document.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
                if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                process_inline(li, p)
        
        elif tag == 'table':
            rows = element.xpath('.//tr')
            if not rows: continue
            cells_in_first_row = rows[0].xpath('./td|./th')
            cols_count = len(cells_in_first_row)
            if cols_count == 0: continue
            
            table = document.add_table(rows=0, cols=cols_count)
            table.style = 'Table Grid'
            for row_el in rows:
                cells = row_el.xpath('./td|./th')
                row = table.add_row()
                for i, cell_el in enumerate(cells):
                    if i < len(row.cells):
                        p = row.cells[i].paragraphs[0]
                        if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        process_inline(cell_el, p)

def process_inline(element, paragraph):
    """Processes text nodes and inline math within an element."""
    inner_html = lxml.html.tostring(element, encoding='unicode', method='html')
    start = inner_html.find('>') + 1
    end = inner_html.rfind('<')
    inner_content = inner_html[start:end] if start < end else element.text_content()

    # Remove <br> tags inside math delimiters so they don't break math parsing
    # Handle display math $$...$$ first, then inline math $...$
    def strip_br_in_math(m):
        return re.sub(r'<br\s*/?>', ' ', m.group(0))
    inner_content = re.sub(r'\$\$[\s\S]*?\$\$', strip_br_in_math, inner_content)
    inner_content = re.sub(r'\$[^\$]+?\$', strip_br_in_math, inner_content)

    # Split into sections of text, display math ($$), and inline math ($)
    parts = re.split(r'(\$\$[\s\S]*?\$\$|\$[^\$]+?\$)', inner_content)
    
    for part in parts:
        if not part: continue
        if part.startswith('$$') and part.endswith('$$'):
            # Unescape LaTeX display math
            latex = html.unescape(part[2:-2])
            add_math_to_run(paragraph, latex, is_display=True)
        elif part.startswith('$') and part.endswith('$'):
            # Unescape LaTeX inline math
            latex = html.unescape(part[1:-1])
            add_math_to_run(paragraph, latex, is_display=False)
        else:
            # Clean text and add to run
            clean_text = re.sub(r'<br\s*/?>', '\n', part)
            clean_text = re.sub(r'<[^>]+>', '', clean_text)
            clean_text = html.unescape(clean_text)
            paragraph.add_run(clean_text)

@app.route('/')
def index():
    visitor_id = request.cookies.get('visitor_id')
    response = make_response(render_template('index.html'))
    
    if not visitor_id:
        visitor_id = str(uuid.uuid4())
        response.set_cookie('visitor_id', visitor_id, max_age=60*60*24*365) # 1 year
        today = datetime.date.today().isoformat()
        try:
            with sqlite3.connect(DB_PATH) as conn:
                c = conn.cursor()
                c.execute('INSERT OR IGNORE INTO users (visitor_id, first_visit_date) VALUES (?, ?)', (visitor_id, today))
                if c.rowcount > 0:
                    c.execute('INSERT OR IGNORE INTO daily_stats (date, downloads, chars, unique_users) VALUES (?, 0, 0, 0)', (today,))
                    c.execute('UPDATE daily_stats SET unique_users = unique_users + 1 WHERE date = ?', (today,))
                conn.commit()
        except sqlite3.Error as e:
            print(f"Database error tracking user: {e}")

    return response

@app.route('/convert', methods=['POST'])
def convert():
    content = request.form.get('content', '')
    is_rtl = request.form.get('is_rtl') == 'on'
    
    # 1. Preprocess
    clean_text = preprocess_copied_math(content)
    
    # 2. Convert to HTML via Markdown
    html_body = markdown.markdown(clean_text, extensions=['extra', 'nl2br'])
    
    # 3. Create Document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 4. Map HTML to Docx
    html_to_docx(f"<body>{html_body}</body>", doc, is_rtl=is_rtl)
    
    # 5. Save and Return
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # 6. Update Stats and History
    visitor_id = request.cookies.get('visitor_id')
    is_new_visitor = False
    if not visitor_id:
        visitor_id = str(uuid.uuid4())
        is_new_visitor = True
    
    try:
        chars_count = len(content)
        today = datetime.date.today().isoformat()
        with sqlite3.connect(DB_PATH) as conn:
            c = conn.cursor()
            # Ensure today exists
            c.execute('INSERT OR IGNORE INTO daily_stats (date, downloads, chars, unique_users) VALUES (?, 0, 0, 0)', (today,))
            
            # Update base stats
            c.execute('UPDATE daily_stats SET downloads = downloads + 1, chars = chars + ? WHERE date = ?', (chars_count, today))
            
            # Log full conversion history
            c.execute('''INSERT INTO conversions (visitor_id, content, char_count, is_rtl) 
                         VALUES (?, ?, ?, ?)''', (visitor_id, content, chars_count, 1 if is_rtl else 0))
            
            # Handle unique users
            # Logic: Check if this user (visitor_id) has converted anything TODAY yet
            c.execute('SELECT id FROM conversions WHERE visitor_id = ? AND date(timestamp) = ? LIMIT 1', (visitor_id, today))
            has_today = c.fetchone()
            
            # If this is their first conversion of the day (including the one we just inserted, 
            # so we check if there's only 1 or if we just inserted the first one), increment unique_users.
            # Actually, since we already inserted, we check if count == 1 for today.
            c.execute('SELECT COUNT(*) FROM conversions WHERE visitor_id = ? AND date(timestamp) = ?', (visitor_id, today))
            if c.fetchone()[0] == 1:
                 c.execute('UPDATE daily_stats SET unique_users = unique_users + 1 WHERE date = ?', (today,))
            
            conn.commit()
    except sqlite3.Error as e:
        print(f"Database error updating stats: {e}")
    
    response = send_file(
        file_stream,
        as_attachment=True,
        download_name="converted.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Set persistent cookie for 1 year
    response.set_cookie('visitor_id', visitor_id, max_age=31536000)
    return response

def check_auth(username, password):
    return username == 'admin' and password == '123456'

def authenticate():
    return Response(
    'Could not verify your access level for that URL.\n'
    'You have to login with proper credentials', 401,
    {'WWW-Authenticate': 'Basic realm="Admin Login Required"'})

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

@app.route('/admin')
@requires_auth
def admin():
    try:
        with sqlite3.connect(DB_PATH) as conn:
            c = conn.cursor()
            # Global totals
            c.execute('SELECT SUM(unique_users), SUM(downloads), SUM(chars) FROM daily_stats')
            totals = c.fetchone()
            if not totals or totals[0] is None:
                totals = (0, 0, 0)
                
            # History Charts
            c.execute('SELECT date, unique_users, downloads FROM daily_stats ORDER BY date DESC LIMIT 14')
            history = [{'date': row[0], 'users': row[1], 'downloads': row[2]} for row in c.fetchall()]
            history.reverse()

            # Detailed conversion logs for the manager
            c.execute('SELECT id, timestamp, content, char_count, is_rtl FROM conversions ORDER BY timestamp DESC LIMIT 50')
            logs = []
            for row in c.fetchall():
                # Strip HTML but keep it readable for preview
                preview = row[2][:100] + '...' if len(row[2]) > 100 else row[2]
                logs.append({
                    'id': row[0],
                    'time': row[1],
                    'content': row[2],
                    'preview': preview,
                    'chars': row[3],
                    'rtl': row[4]
                })
    except sqlite3.Error:
        totals = (0, 0, 0)
        history = []
        logs = []
        
    return render_template('admin.html', stats={
        'users': totals[0],
        'downloads': totals[1],
        'chars': totals[2],
        'history': history,
        'logs': logs,
        'today': datetime.date.today().isoformat()
    })

if __name__ == '__main__':
    app.run(debug=True)
