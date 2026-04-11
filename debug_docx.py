import lxml.html
import markdown
from app import preprocess_copied_math, html_to_docx
from docx import Document

content = "This is a test paragraph.\n\n$$E = mc^2$$"
clean_text = preprocess_copied_math(content)
html_body = markdown.markdown(clean_text, extensions=['extra', 'nl2br'])
full_html = f"<body>{html_body}</body>"

print(f"--- FULL HTML ---\n{full_html}\n")

soup = lxml.html.fromstring(full_html)
elements = soup.xpath('body/*')
print(f"Found {len(elements)} elements in body.")

for el in elements:
    print(f"Tag: {el.tag}, Text: {el.text_content()}")

doc = Document()
html_to_docx(full_html, doc)
print(f"Doc paragraphs count: {len(doc.paragraphs)}")

for p in doc.paragraphs:
    print(f"P Text: {p.text}")
