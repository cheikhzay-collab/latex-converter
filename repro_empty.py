import io
import re
import markdown
import latex2mathml.converter
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from lxml import etree
import lxml.html

# Copying functions from app.py to test
def preprocess_copied_math(text):
    text = re.sub(r'\\\[([\s\S]*?)\\\]', r'$$\1$$', text)
    def bracket_math_cleaner(match):
        content = match.group(1).strip()
        if re.search(r'^[\d., \-+]+[ \t]*[,;][ \t]*[\d., \-+\\infty]+[\[\]]?$', content):
            return match.group(0) # Keep as is
        if re.search(r'[\^\\=_]|\d+[a-z]|\\frac', content) or len(content) > 15:
            return f"$${content}$$"
        return match.group(0)

    text = re.sub(r'(?m)^[ \t]*\[([^\[\]\n]+)\][ \t]*$', bracket_math_cleaner, text)
    text = re.sub(r'\(\((.*?)\)\)', r'\(\1\)', text)
    text = re.sub(r'\[cite[_:][^\]]*\]', '', text)
    return text

def html_to_docx(html_content, document, is_rtl=False):
    if not html_content.strip():
        return
        
    try:
        soup = lxml.html.fromstring(html_content)
    except Exception as e:
        print(f"HTML Parse Error: {e}")
        return

    elements = soup.xpath('//body/*') or soup.xpath('/*')
    print(f"Found {len(elements)} elements")
    
    for element in elements:
        tag = element.tag
        print(f"Processing tag: {tag}")
        if tag == 'p':
            p = document.add_paragraph()
            print("Added paragraph")
            process_inline(element, p)

def process_inline(element, paragraph):
    inner_html = lxml.html.tostring(element, encoding='unicode', method='html')
    start = inner_html.find('>') + 1
    end = inner_html.rfind('<')
    inner_content = inner_html[start:end] if start < end else element.text_content()
    print(f"Inner content: {inner_content}")

    parts = re.split(r'(\$\$[\s\S]*?\$\$|\$[^\$\n]+?\$)', inner_content)
    
    for part in parts:
        if not part: continue
        print(f"Part: {part}")
        paragraph.add_run(part)

# Test case
content = "Hello world"
clean_text = preprocess_copied_math(content)
html_body = markdown.markdown(clean_text, extensions=['extra', 'nl2br'])
print(f"HTML Body: {html_body}")

doc = Document()
html_to_docx(f"<body>{html_body}</body>", doc)

doc.save("test_output.docx")
print("Saved test_output.docx")
