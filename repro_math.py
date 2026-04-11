from docx import Document
from docx.oxml import OxmlElement
import latex2mathml.converter

def create_omath_element(mathml_string):
    # This is where the magic needs to happen.
    # Inserting raw MathML into python-docx is not directly supported.
    # We need to construct the OMML structure.
    
    # Since we don't have a MathML -> OMML converter, 
    # we might need to rely on the fact that Word can interpret MathML
    # IF it's inserted as an ``<m:oMathPara>`` with the right namespace.
    
    # However, 'latex2mathml' produces standard MathML.
    # Word requires a specific wrapper or transformation (MML2OMML.XSL).
    
    # Let's try to just insert the MathML as a string and see if we can trick Word
    # into displaying it? Unlikely.
    
    # The robust way in Python without external binaries (like Pandoc) is surprisingly hard.
    # BUT, we can use the `lxml` library to build the XML tree if we know the structure.
    
    # Wait, there IS a way to insert MathML directly into the document.xml
    # if we wrap it correctly.
    
    pass

def test_doc_gen():
    doc = Document()
    p = doc.add_paragraph("Test Equation below:")
    
    latex = "x = \\frac{-b \\pm \\sqrt{b^2 - 4ac}}{2a}"
    mathml = latex2mathml.converter.convert(latex)
    
    print("Generated MathML:", mathml)
    
    # Try to add it as a run with the mathml (this usually just shows text)
    # p.add_run(mathml) 
    
    doc.save("test_equation.docx")

if __name__ == "__main__":
    test_doc_gen()
